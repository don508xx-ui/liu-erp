"""生成喷涂加工ERP使用白皮书(Word)到桌面。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os, datetime

OUT = r"C:\Users\Administrator\Desktop\喷涂加工ERP-使用白皮书.docx"

# 配色(深色科技风)
C_PRIMARY = RGBColor(0x00, 0xa8, 0xe8)      # 主青蓝
C_DARK = RGBColor(0x0a, 0x0e, 0x1a)         # 深底
C_PANEL = RGBColor(0x1a, 0x21, 0x38)        # 面板
C_TEXT = RGBColor(0x22, 0x2a, 0x3a)         # 正文深灰蓝
C_TEXT2 = RGBColor(0x5a, 0x66, 0x80)        # 次级
C_ACCENT = RGBColor(0x7c, 0x3a, 0xed)       # 紫
C_SUCCESS = RGBColor(0x10, 0xb9, 0x81)
C_WHITE = RGBColor(0xff, 0xff, 0xff)

FONT_CN = "微软雅黑"
FONT_EN = "Segoe UI"

doc = Document()

# ===== 全局样式 =====
def set_run_font(run, size=10.5, bold=False, color=C_TEXT, font_cn=FONT_CN):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = FONT_EN
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), font_cn)

# 默认正文样式
style = doc.styles['Normal']
style.font.name = FONT_EN
style.font.size = Pt(10.5)
style.font.color.rgb = C_TEXT
style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# 页面边距
for sec in doc.sections:
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.4)


def shade(cell, hex_color):
    """单元格底色"""
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hex_color)
    tcPr.append(sh)


def set_cell_borders(cell, color="d0d7e2", sz=6):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)


def add_heading(text, level=1, color=None):
    """章节标题"""
    color = color or C_PRIMARY
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, size=18, bold=True, color=C_DARK)
        # 左侧色条
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '36')
        left.set(qn('w:space'), '8')
        left.set(qn('w:color'), '00a8e8')
        pbdr.append(left)
        pPr.append(pbdr)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, size=13.5, bold=True, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11.5, bold=True, color=C_TEXT)
    return p


def add_para(text, size=10.5, color=C_TEXT, bold=False, indent=0, align=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(text, size=10.5, color=C_TEXT):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.4
    run = p.runs[0] if p.runs else p.add_run('')
    run.text = text
    set_run_font(run, size=size, color=color)
    return p


def add_table(headers, rows, col_widths=None, header_fill='1a2138', zebra='f3f6fb'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # 表头
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=C_WHITE)
        shade(hdr[i], header_fill)
        set_cell_borders(hdr[i], color="1a2138", sz=4)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 数据行
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            p.paragraph_format.line_spacing = 1.3
            run = p.add_run(str(val))
            set_run_font(run, size=9.5, color=C_TEXT)
            if ri % 2 == 1:
                shade(cells[ci], zebra)
            set_cell_borders(cells[ci])
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_callout(title, text, fill='eaf6fd', border='00a8e8'):
    """提示框"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = ''
    shade(cell, fill)
    set_cell_borders(cell, color=border, sz=12)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    set_run_font(r1, size=10.5, bold=True, color=C_PRIMARY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.4
    r2 = p2.add_run(text)
    set_run_font(r2, size=10, color=C_TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def page_break():
    doc.add_page_break()


# ===== 页眉页脚 =====
def setup_header_footer():
    sec = doc.sections[0]
    # 页眉
    hdr = sec.header
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("喷涂加工ERP · 业财一体化  |  使用白皮书")
    set_run_font(hr, size=8.5, color=C_TEXT2)
    # 页脚(页码)
    ftr = sec.footer
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    set_run_font(run, size=9, color=C_TEXT2)
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldEnd)
    prefix = fp.add_run("  /  ")
    set_run_font(prefix, size=9, color=C_TEXT2)
    run2 = fp.add_run()
    set_run_font(run2, size=9, color=C_TEXT2)
    b2 = OxmlElement('w:fldChar'); b2.set(qn('w:fldCharType'), 'begin')
    i2 = OxmlElement('w:instrText'); i2.set(qn('xml:space'), 'preserve'); i2.text = 'NUMPAGES'
    e2 = OxmlElement('w:fldChar'); e2.set(qn('w:fldCharType'), 'end')
    run2._r.append(b2); run2._r.append(i2); run2._r.append(e2)


setup_header_footer()

# ==================== 封面 ====================
# 顶部留白
for _ in range(3):
    doc.add_paragraph()

# 深色背景标题块
cover_tbl = doc.add_table(rows=1, cols=1)
cell = cover_tbl.rows[0].cells[0]
shade(cell, "0a0e1a")
set_cell_borders(cell, color="0a0e1a", sz=0)
cell.text = ''
# 主标题
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
p.paragraph_format.space_after = Pt(8)
r = p.add_run("喷涂加工企业 ERP")
set_run_font(r, size=32, bold=True, color=C_WHITE)
# 副标题
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
r2 = p2.add_run("业财一体化 · 智能运营平台")
set_run_font(r2, size=15, bold=False, color=C_PRIMARY)
# 装饰线
p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(24)
r3 = p3.add_run("━━━━━━━━━━━━━━━━━━━━━━━━")
set_run_font(r3, size=10, color=C_PRIMARY)
# 英文
p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(36)
r4 = p4.add_run("Spray Processing ERP  ·  Business-Finance Integration Platform")
set_run_font(r4, size=10, color=RGBColor(0x8a, 0x96, 0xb3), font_cn=FONT_EN)

# 封面信息
info_tbl = doc.add_table(rows=1, cols=1)
icell = info_tbl.rows[0].cells[0]
icell.text = ''
shade(icell, "f3f6fb")
set_cell_borders(icell, color="00a8e8", sz=8)
items = [
    ("文档类型", "产品使用白皮书 · 客户体验 DEMO"),
    ("适用版本", "V1.0.0"),
    ("发布日期", datetime.date.today().strftime("%Y-%m-%d")),
    ("文档密级", "公开 · 可对外分发"),
]
first = icell.paragraphs[0]
first.paragraph_format.space_before = Pt(10)
for i, (k, v) in enumerate(items):
    p = icell.paragraphs[0] if i == 0 else icell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    rk = p.add_run(f"  {k}：  ")
    set_run_font(rk, size=10.5, bold=True, color=C_TEXT2)
    rv = p.add_run(v)
    set_run_font(rv, size=10.5, bold=True, color=C_DARK)
icell.paragraphs[-1].paragraph_format.space_after = Pt(10)

# 封底说明
for _ in range(2):
    doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = note.add_run("本白皮书配套在线 DEMO 体验环境，具体访问方式见第二章。")
set_run_font(nr, size=10, color=C_TEXT2)

page_break()

# ==================== 第一章 系统概述 ====================
add_heading("第一章  系统概述", level=1)

add_heading("1.1  产品定位", level=2)
add_para("本系统是面向表面喷涂加工企业的一体化 ERP 平台，覆盖销售接单、生产派工、物料领用、"
         "完工成本归集、财务核算、采购供应链、薪酬审批等全业务链路。核心解决喷涂加工企业长期存在的"
         "「业务与财务脱节」痛点——业务发生的瞬间，财务单据与库存流水自动生成，确保业务、库存、资金三本账实时一致。")

add_heading("1.2  核心特色", level=2)
features = [
    ("业财联动", "事件总线 + 钩子驱动。订单生效、领料确认、完工确认等关键业务节点，自动生成应收应付、库存流水、成本归集凭证，零手工传递。"),
    ("工单成本归集", "所有材料、人工、制造费用强制挂工单。完工确认瞬间实时汇总工单总成本、订单收入与毛利，盈亏一目了然。"),
    ("喷涂行业特化", "客供料台账、涂料利用率追踪（理论用量 vs 实际用量）、批次追溯、返工成本独立核算，贴合喷涂场景。"),
    ("表驱动审批", "审批流程配置化，运营后台自助调整节点与角色，无需改代码。"),
    ("Agent 接口预留", "内置 scoped token 机制的 Agent API，支持未来接入智能体做自动查询、预警触发、报表生成。"),
    ("开箱即用", "本地化部署，安装即用，数据不上云，贴合中小喷涂企业 IT 现状。"),
]
add_table(["特色能力", "说明"], features, col_widths=[3.5, 12.5])

add_heading("1.3  技术架构", level=2)
arch = [
    ("后端", "Python 3.10 · FastAPI · SQLAlchemy 2.0 · JWT 鉴权 · 事件总线 · SQLite(可平滑迁移 PostgreSQL)"),
    ("前端", "Vue 3 · Element Plus · ECharts · 深色科技风 UI · 响应式布局"),
    ("部署", "Uvicorn 应用服务器 · Docker 容器化(可选) · 局域网/公网双模访问"),
    ("集成", "企业微信接口预留 · Agent API · 邮件/语音通知栈"),
]
add_table(["层级", "技术栈"], arch, col_widths=[2.5, 13.5])

page_break()

# ==================== 第二章 DEMO访问 ====================
add_heading("第二章  DEMO 体验访问", level=1)

add_heading("2.1  访问方式", level=2)
add_callout("🌐 在线 DEMO 地址",
            "https://stanford-travels-advisor-candles.trycloudflare.com",
            fill='eaf6fd', border='00a8e8')
add_para("在浏览器地址栏输入上述 URL 即可打开系统登录页。推荐使用 Chrome、Edge 等现代浏览器，"
         "无需安装任何客户端插件。", space_after=8)

add_heading("2.2  登录账号", level=2)
accounts = [
    ("admin", "123456", "系统管理员", "全部模块 · 全部权限(体验首选)"),
    ("sales01", "123456", "销售", "客户、订单、应收"),
    ("ops01", "123456", "运营助理", "加工单、完工、库存"),
    ("fin01", "123456", "财务", "财务、收款、薪酬"),
    ("wh01", "123456", "仓管", "库存、领料、采购"),
    ("gm01", "123456", "总经理", "全模块只读 · 审批"),
]
add_table(["用户名", "密码", "角色", "可操作范围"], accounts, col_widths=[2.8, 2.2, 2.5, 8.0])

add_heading("2.3  浏览器要求", level=2)
add_bullet("Chrome 90+ 或 Microsoft Edge 90+（推荐）")
add_bullet("Firefox 88+ 或 Safari 14+")
add_bullet("开启 JavaScript，允许浏览器存储 localStorage（登录态依赖）")
add_bullet("屏幕分辨率建议 1366×768 及以上，移动端可访问但为桌面布局")

add_heading("2.4  重要提醒", level=2)
add_callout("⚠ 体验须知",
            "1. DEMO 数据为模拟数据，请勿录入真实客户或商业信息；\n"
            "2. 体验环境通过临时隧道提供，URL 非永久有效，如无法访问请联系提供方刷新；\n"
            "3. 系统每日会重置示例数据，您操作产生的变更不会被保留；\n"
            "4. 多人可同时登录体验，数据实时联动，您可以看到他人操作的留痕。",
            fill='fff7e6', border='f59e0b')

page_break()

# ==================== 第三章 业务流程 ====================
add_heading("第三章  业务流程", level=1)

add_heading("3.1  自营料订单全流程", level=2)
add_para("客户自带图纸、由我方采购涂料进行加工的标准模式。业务链路如下：", space_after=6)
flow_a = [
    ("1", "销售建客户", "客户模块录入客户资料（编码/名称/税号/结算周期等），客户ID是全系统关联枢纽"),
    ("2", "销售下订单", "订单模块新建订单，关联客户，录入零件、数量、单价、涂料规格、加工要求"),
    ("3", "订单生效", "订单提交→生效，钩子自动生成预收款应收凭证"),
    ("4", "运营下加工单", "订单生效后可下加工单，指定车间、批次、计划数量"),
    ("5", "加工单下达", "下达后钩子自动生成领料单（按BOM算涂料理论用量）"),
    ("6", "仓管确认领料", "确认领料→扣减库存→记录材料成本到工单"),
    ("7", "完工填报", "录入完工数/合格数/返工数/报废数、人工工时与费用"),
    ("8", "完工确认", "钩子归集材料+人工+制费，算出工单总成本与订单毛利"),
    ("9", "财务登记收款", "登记收款单→自动核销应收账款"),
]
add_table(["步骤", "环节", "说明"], flow_a, col_widths=[1.5, 3.0, 11.5])

add_heading("3.2  客供料订单全流程", level=2)
add_para("客户提供物料（来料加工）模式，系统自动走客供料台账，不扣本方库存：", space_after=6)
flow_b = [
    ("1", "建客户+下订单", "同自营料，但订单明细「物料模式」选「客供料」"),
    ("2", "订单生效+下加工单", "同自营料"),
    ("3", "加工单下达", "钩子识别客供料，不生成领料单，自动记客供料台账（来料验收记录）"),
    ("4", "完工确认", "无材料成本，只归集人工与制费"),
    ("5", "收款核销", "同自营料"),
]
add_table(["步骤", "环节", "说明"], flow_b, col_widths=[1.5, 3.5, 11.0])

add_heading("3.3  业财联动机制", level=2)
add_para("以下业务动作会自动触发财务/库存凭证，无需手工传递单据：", space_after=6)
hooks = [
    ("订单生效", "order.effective", "生成预收款应收凭证（按预收比例）"),
    ("加工单下达", "work_order.released", "自动生成领料单 + 通知车间厂长"),
    ("领料确认", "material.confirmed", "扣减库存 + 记录材料成本到工单"),
    ("完工确认", "completion.confirmed", "归集人工/制费 + 算工单总成本 + 记库存入库 + 算利润"),
    ("收款登记", "receipt.created", "自动核销应收账款 + 触发收款通知"),
]
add_table(["业务动作", "事件标识", "自动触发的财务/库存动作"], hooks, col_widths=[2.8, 3.8, 9.4])

page_break()

# ==================== 第四章 模块操作 ====================
add_heading("第四章  功能模块操作指南", level=1)
add_para("系统共 17 个业务模块，左侧菜单切换。每个模块支持查询、新增、编辑、详情等操作。"
         "下表为各模块速查，后续按模块说明。", space_after=8)

modules = [
    ("工作台", "KPI看板、成本趋势、应收账龄、待办预警"),
    ("客户管理", "客户档案、结算周期、开户行信息"),
    ("订单管理", "销售订单全生命周期（草稿→生效→关闭）"),
    ("加工单管理", "工单创建、下达、车间派工"),
    ("完工管理", "完工填报、合格/返工/报废、成本归集"),
    ("领料管理", "自动领料单、库存校验、确认/驳回"),
    ("库存管理", "物料主数据、实时库存、安全库存"),
    ("财务管理", "应收应付、收款登记、凭证查询"),
    ("采购管理", "采购单、入库、应付"),
    ("采购申请", "各部门请购、汇总转采购"),
    ("薪酬管理", "工资单、人工成本归集"),
    ("审批中心", "配置化审批流、待办处理"),
    ("预警中心", "规则引擎、库存/应收/异常预警"),
    ("通知中心", "站内信、多渠道通知"),
    ("库存流水", "出入库明细、批次追溯"),
    ("客供料台账", "客供料来料/消耗/结存"),
    ("Agent接口", "Scoped Token、API对接"),
]
add_table(["模块", "核心功能"], modules, col_widths=[3.0, 13.0])

# 各模块详细说明
add_heading("4.1  工作台", level=2)
add_bullet("顶部 KPI 卡片：本月订单额、完工产值、应收余额、待处理预警数")
add_bullet("成本趋势图：近 30 天工单成本走势（材料/人工/制费分项）")
add_bullet("应收账龄图：0-30/30-60/60-90/90+ 天应收分布")
add_bullet("进入工作台自动检查预警规则，红点提示待处理项")

add_heading("4.2  客户管理", level=2)
add_bullet("「+ 新增客户」录入客户编码、名称、税号、联系人、结算周期、开户行等")
add_bullet("客户ID是订单、加工单、完工、收款、客供料台账、应收的共同关联键")
add_bullet("支持按名称/编码搜索，点击「详情」查看客户完整档案")
add_bullet("结算周期字段：月结30/60/90/款到发货，影响应收账龄统计")

add_heading("4.3  订单管理", level=2)
add_bullet("「+ 新建订单」选择客户，录入零件明细（名称/规格/计价方式/数量/单价）")
add_bullet("物料模式：自营料（默认）/ 客供料，决定后续是否生成领料单")
add_bullet("订单状态：草稿→提交→生效→关闭，生效后才可下加工单")
add_bullet("支持退单（记录退单次数与原因）、查看订单利润")

add_heading("4.4  加工单管理", level=2)
add_bullet("订单生效后，「+ 新建加工单」关联订单，指定车间(A/B)、批次号、计划数量")
add_bullet("加工单状态：创建→下达→加工中→完工，下达后自动生成领料单")
add_bullet("支持外协加工（关联供应商、外协费用）")

add_heading("4.5  完工管理", level=2)
add_bullet("「+ 完工填报」录入完工数、合格数、返工数、报废数")
add_bullet("录入人工工时与人工成本、制造费用")
add_bullet("若不填明细，系统自动从已确认领料单带出理论/实际用量")
add_bullet("确认完工后，钩子自动归集成本、算毛利、记库存入库")

add_heading("4.6  领料管理", level=2)
add_bullet("领料单由加工单下达时自动生成，无需手工创建")
add_bullet("仓管「确认」前校验库存是否充足，不足则驳回")
add_bullet("确认后扣减库存、记录材料成本到对应工单")

add_heading("4.7  库存管理", level=2)
add_bullet("物料分类：涂料粉末、溶剂、辅料、客供料等")
add_bullet("实时库存、安全库存阈值，低于阈值触发预警")
add_bullet("「+ 新增物料」录入名称、规格、单位、单价、安全库存")

add_heading("4.8  财务管理", level=2)
add_bullet("应收应付凭证自动生成，支持按类型/状态筛选")
add_bullet("「登记收款」选择订单→录入金额→自动核销应收")
add_bullet("工单成本查询：材料/人工/制费分项明细")
add_bullet("订单利润查询：收入、成本、毛利、毛利率")

add_heading("4.9  采购管理", level=2)
add_bullet("「+ 新建采购单」选择供应商、录入物料与数量")
add_bullet("采购入库后自动增加库存、生成应付凭证")

add_heading("4.10  采购申请", level=2)
add_bullet("各部门发起请购，汇总后可转采购单")

add_heading("4.11  薪酬管理", level=2)
add_bullet("新建工资单，归集人工成本到工单/订单")

add_heading("4.12  审批中心", level=2)
add_bullet("配置化审批流，支持多节点、多角色")
add_bullet("待办列表一键通过/驳回")

add_heading("4.13  预警中心", level=2)
add_bullet("规则引擎：库存不足、应收超期、订单延期等")
add_bullet("「立即检查所有规则」手动触发全量扫描")

add_heading("4.14  通知中心", level=2)
add_bullet("站内信按渠道筛选，记录已读/未读")

add_heading("4.15  库存流水", level=2)
add_bullet("出入库明细，按类型（出库/入库/退料/采购）筛选")
add_bullet("批次追溯：通过批次号查全链路流转")

add_heading("4.16  客供料台账", level=2)
add_bullet("客供料来料验收、消耗、结存记录")
add_bullet("与客供料订单自动关联")

add_heading("4.17  Agent 接口", level=2)
add_bullet("「生成测试 Token」创建 scoped API Token")
add_bullet("支持未来接入智能体做自动查询、预警、报表生成")
add_bullet("Token 限定读写权限，保障系统安全")

page_break()

# ==================== 第五章 角色权限 ====================
add_heading("第五章  角色与权限", level=1)
add_para("系统采用 RBAC 角色权限模型，不同角色看到和能操作的范围不同。"
         "体验时建议先用 admin 账号走全流程，再切换其他角色感受权限差异。", space_after=8)
roles = [
    ("ADMIN", "系统管理员", "全部模块全部操作（体验首选）"),
    ("SALES", "销售", "客户、订单增删改，查看本人订单财务"),
    ("OPERATION", "运营助理", "加工单、完工、库存读取"),
    ("FINANCE", "财务", "财务、收款、薪酬、采购只读"),
    ("WAREHOUSE", "仓管", "库存、领料、采购"),
    ("GM", "总经理", "全模块只读，审批"),
    ("MANAGER", "车间厂长", "本车间工单"),
]
add_table(["角色代码", "角色名", "权限范围"], roles, col_widths=[2.8, 2.8, 10.4])

page_break()

# ==================== 第六章 体验路径建议 ====================
add_heading("第六章  推荐体验路径", level=1)
add_para("为帮助您快速感受系统价值，建议按以下路径体验（约 15 分钟）：", space_after=8)

add_heading("路径一：感受业财联动（推荐）", level=2)
steps1 = [
    ("1", "admin 登录，进入「工作台」查看 KPI 与图表"),
    ("2", "进入「客户」→「+ 新增客户」，任意填一个客户"),
    ("3", "进入「订单」→「+ 新建订单」，关联刚建的客户，录入1行零件，保存"),
    ("4", "订单列表点「提交」「生效」，进入「财务」查看应收凭证已自动生成"),
    ("5", "进入「加工单」→「+ 新建加工单」关联订单，保存后点「下达」"),
    ("6", "进入「领料」，看到自动生成的领料单，点「确认」"),
    ("7", "进入「完工」→「+ 完工填报」，录入完工数，点「确认」"),
    ("8", "回到「财务」→ 工单成本 / 订单利润，查看自动归集的成本与毛利"),
    ("9", "「登记收款」，查看应收被自动核销"),
]
for s, t in steps1:
    add_bullet(f"【步骤{s}】{t}")

add_heading("路径二：感受客供料模式", level=2)
steps2 = [
    ("1", "新建订单时，零件明细「物料模式」选「客供料」"),
    ("2", "订单生效→下加工单→下达，进入「客供料」模块查看自动记账"),
    ("3", "领料模块不会生成领料单（料是客户的），完工只归集人工费"),
]
for s, t in steps2:
    add_bullet(f"【步骤{s}】{t}")

add_heading("路径三：感受喷涂行业特化", level=2)
add_bullet("完工填报时查看「理论用量 vs 实际用量」，体现涂料利用率追踪")
add_bullet("「库存流水」按批次号追溯涂料流转")
add_bullet("「客供料台账」单独管理来料加工的物料")

page_break()

# ==================== 第七章 FAQ ====================
add_heading("第七章  常见问题", level=1)

faqs = [
    ("Q1：DEMO 地址打不开？",
     "临时隧道偶有延迟，请刷新或过几分钟再试；如持续不可达，请联系提供方刷新隧道。"),
    ("Q2：登录提示无角色/无权限？",
     "请确认使用白皮书第二章列出的账号（admin/sales01 等），密码统一为 123456。"),
    ("Q3：为什么我新增的数据不见了？",
     "DEMO 环境数据每日会重置，体验产生的变更为演示性质，不持久保留。"),
    ("Q4：可以多人同时体验吗？",
     "可以，多人同时登录操作，数据实时联动，您能看到他人操作的留痕。"),
    ("Q5：手机能访问吗？",
     "可以，但当前为桌面布局，手机浏览会有挤压，建议 PC 体验为佳。"),
    ("Q6：如何接入我们自己的数据？",
     "正式版支持导入 Excel 客户/物料/期初库存，详情咨询提供方。"),
    ("Q7：系统能对接企业微信吗？",
     "已预留企业微信接口，正式版可配置通知推送、单据审批到企微。"),
    ("Q8：Agent 接口是干什么的？",
     "为未来接入 AI 智能体预留，可做自然语言查询、自动预警、智能报表等扩展能力。"),
]
for q, a in faqs:
    add_heading(q, level=3, color=C_ACCENT)
    add_para(a, indent=0.4, color=C_TEXT2)

# ==================== 结尾 ====================
page_break()
add_heading("联系我们", level=1)
add_para("如需正式版部署、定制开发、数据导入或培训支持，请联系提供方。", space_after=10)
contact = [
    ("体验支持", "扫码或邮件反馈体验问题"),
    ("商务咨询", "正式版授权与部署方案"),
    ("技术对接", "API 集成、企业微信、Agent 接入"),
]
add_table(["事项", "说明"], contact, col_widths=[3.0, 13.0])

end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
end.paragraph_format.space_before = Pt(40)
er = end.add_run("— 本白皮书至此结束 · 感谢您的体验 —")
set_run_font(er, size=10, color=C_TEXT2)

# 保存
doc.save(OUT)
print(f"✅ 已生成: {OUT}")
print(f"   大小: {os.path.getsize(OUT)/1024:.1f} KB")

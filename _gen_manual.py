# -*- coding: utf-8 -*-
"""生成《峰业精密ERP 系统功能及详细操作手册》Word文档"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============ 常量 ============
PRIMARY = RGBColor(0x1F, 0x3A, 0x6E)   # 深蓝
ACCENT  = RGBColor(0x0E, 0x7C, 0x86)   # 青绿
GRAY    = RGBColor(0x66, 0x66, 0x66)
HEADER_BG = "1F3A6E"
ALT_BG    = "F2F6FC"
FONT = "微软雅黑"

doc = Document()

# ============ 全局样式 ============
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
style.paragraph_format.line_spacing = 1.4
style.paragraph_format.space_after = Pt(4)

for sec in doc.sections:
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)


def set_cn(run, font=FONT, size=None, bold=None, color=None, italic=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color: run.font.color.rgb = color
    run.font.italic = italic


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def heading(text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    color = PRIMARY if level == 1 else ACCENT
    set_cn(r, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True, color=color)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(6)
    return h


def para(text, size=10.5, bold=False, color=None, indent=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    set_cn(r, size=size, bold=bold, color=color)
    return p


def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_cn(r1, bold=True)
    r = p.add_run(text)
    set_cn(r)
    return p


def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_cn(r1, bold=True)
    r = p.add_run(text)
    set_cn(r)
    return p


def make_table(headers, rows, widths=None, header_bg=HEADER_BG, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_cn(r, size=font_size, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(hdr[i], header_bg)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            set_cn(r, size=font_size)
            if ri % 2 == 1:
                shade_cell(cells[ci], ALT_BG)
    if widths:
        for ci, w in enumerate(widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def flow_box(title, steps):
    """流程节点步骤图框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("▶ " + title)
    set_cn(r, size=10, bold=True, color=ACCENT)
    pp = doc.add_paragraph()
    pp.paragraph_format.left_indent = Cm(0.4)
    rr = pp.add_run("  ➤  ".join(steps))
    set_cn(rr, size=9.5, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def note(text, kind="info"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    mark = {"info": "💡 提示", "warn": "⚠️ 注意", "danger": "🔴 重要", "ok": "✅ 完成"}.get(kind, "💡 提示")
    r1 = p.add_run(mark + "：")
    set_cn(r1, size=10, bold=True, color=ACCENT if kind in ("info", "ok") else (RGBColor(0xC0, 0x6A, 0x1F) if kind == "warn" else RGBColor(0xC0, 0x3A, 0x2B)))
    r2 = p.add_run(text)
    set_cn(r2, size=10)
    return p


# ================= 封面 =================
for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("峰业精密ERP")
set_cn(r, size=32, bold=True, color=PRIMARY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("系统功能及详细操作手册")
set_cn(r, size=22, bold=True, color=ACCENT)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Vue3 + FastAPI + SQLite  ·  喷涂加工 · 业财一体化  ·  v1.0")
set_cn(r, size=12, color=GRAY)
for _ in range(8):
    doc.add_paragraph()

doc.add_page_break()

# ================= 目录说明 =================
heading("目录", 1)
toc_items = [
    "1  系统概述", "2  登录与系统入口", "3  角色与权限说明", "4  工作台使用说明",
    "5  业务模块详细操作", "6  业务流程操作指引", "7  审批中心", "8  经营分析与AI助手",
    "9  流程设计器", "10  管理员后台", "11  生产大屏", "12  常见问题", "13  附录",
]
for it in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(it)
    set_cn(r, size=11)

doc.add_page_break()

# ================= 1 系统概述 =================
heading("1  系统概述", 1)
para("峰业精密ERP 是一套面向金属表面处理（喷涂加工）行业的一体化经营管理平台，覆盖「销售 → 生产 → 采购 → 仓储 → 财务 → 审批 → 分析」全链路。系统采用角色权限隔离，不同岗位登录后仅能看到与自己业务相关的菜单、工作流与数据。")
make_table(["维度", "说明"], [
    ["技术栈", "前端 Vue3 + Element Plus；后端 FastAPI (Python)；数据库 SQLite (WAL模式)；部署支持 Docker / Zeabur"],
    ["业务覆盖", "客户档案、销售订单、调价申请、来货登记、加工工单、完工单、领料出库、库存、采购申请/订单、财务单据、工资、费用报销、审批中心"],
    ["流程引擎", "内置 7 条标准业务流程（核心生产流、来货登记、费用报销、采购请求、调价申请、采购审批、完工确认），支持可视化流程设计器自定义"],
    ["智能分析", "DeepSeek 双模型（意图解析 + 报告生成），支持自然语言问答与多维透视分析"],
    ["角色体系", "系统管理员、总经理、销售、运营助理、财务、仓管、车间厂长、部门主管、采购 共 9 类角色"],
], widths=[3.5, 12.0])

# ================= 2 登录 =================
heading("2  登录与系统入口", 1)
heading("2.1  访问地址", 2)
make_table(["场景", "地址", "备注"], [
    ["本地开发", "http://localhost:8000", "启动：uvicorn app.main:app --host 0.0.0.0 --port 8000"],
    ["生产环境", "*.zeabur.app 自定义域名", "需正确配置环境变量与健康检查"],
], widths=[3.0, 5.5, 7.0])
heading("2.2  默认账号", 2)
make_table(["账号", "姓名", "角色", "密码"], [
    ["admin", "系统管理员", "系统管理员（全权限）", "admin123"],
    ["sales01", "张销售", "销售", "123456"],
    ["ops01", "王运营", "运营助理", "123456"],
    ["fin01", "赵财务", "财务", "123456"],
    ["wh01", "钱仓管", "仓管", "123456"],
    ["gm01", "孙总/刘总", "总经理", "123456"],
    ["mgr_a / mgr_b", "周厂长A / 吴厂长B", "车间厂长", "123456"],
    ["dept01", "李主管", "部门主管", "123456"],
    ["purchase01", "采购小王", "采购", "123456"],
], widths=[4.0, 3.5, 4.0, 2.5])
note("首次访问新域名时，因浏览器缓存无登录态可能提示“登录已过期”，直接输入账号密码登录即可，无需其他操作。")

# ================= 3 角色权限 =================
heading("3  角色与权限说明", 1)
para("系统严格遵循“角色只能看见与自己相关的工作流”原则，跨角色数据不可见。下表汇总各角色的菜单模块、可见工作流与独有能力。")
make_table(["角色", "可见菜单/模块", "可见工作流（审批链）", "独有能力"], [
    ["ADMIN 系统管理员", "全部模块 + 用户管理 + 角色管理 + 流程设计", "全部 7 条", "增删用户/角色、编辑/删除流程定义、重置密码"],
    ["GM 总经理", "订单/工单/库存/财务/工资/客户/审批/AI分析", "费用报销终审、调价审批、采购审批、核心生产流、完工确认、来货登记", "5000元以上支出终审；AI数据分析权限；查看全部数据"],
    ["SALES 销售", "客户/订单/调价/审批/待办/已办/AI", "调价申请(发起)、来货登记(发起)、核心生产流(发起)", "客户与合同全生命周期；核心生产流业务发起人"],
    ["OPERATION 运营助理", "订单/工单/完工/领料/审批中心/客户", "来货登记运营核对、完工单运营归档、核心生产流", "运营审核入口"],
    ["FINANCE 财务", "财务/工资/采购/库存/审批/报销", "来货登记财务入账、费用报销财务审核、采购财务审核、核心生产流、完工确认", "收款分“一般纳税人(增票)”与“小规模(普票/现金)”两套账"],
    ["WAREHOUSE 仓管", "库存/领料/采购/来货", "核心生产流、来货登记", "来货登记、领料收发存"],
    ["MANAGER 车间厂长", "工单/完工/领料/生产大屏", "完工单提交+质检确认、核心生产流", "车间生产进度大屏"],
    ["DEPARTMENT_HEAD 部门主管", "采购请求/费用报销/审批中心", "费用报销初审、采购请求审批、核心生产流", "5000元以下支出初审"],
    ["PURCHASE 采购", "采购/库存/请求", "采购相关流程", "采购执行、询价"],
], widths=[3.2, 4.6, 4.6, 3.6], font_size=8.5)
note("非管理员访问 /api/admin/* 管理接口会立即返回 403；前端对应入口不会渲染。GM 角色为全权限角色，可见所有导航与数据。")

# ================= 4 工作台 =================
heading("4  工作台使用说明", 1)
para("工作台是登录后的默认首页，所有角色恒有以下四大区域：")
make_table(["区域", "内容", "操作"], [
    ["📈 经营概览", "今日订单、本月销售额、待审批数、低库存告警 KPI 卡片", "点击卡片跳转对应模块"],
    ["📋 我的待办", "当前角色的审批待办与业务待办，带数量徽章（红=紧急/橙=重要/灰=普通）", "点击待办跳转处理页面；“查看全部”进入我的待办列表"],
    ["🔀 业务流程", "当前角色涉及的所有工作流实例，以“节点+箭头”形式平铺展示进度", "点击流程卡片查看节点状态与审批历史"],
    ["📊 工作台", "最近已办（时间轴）+ 团队动态（跨角色协作轨迹）", "查看历史处理记录"],
], widths=[3.2, 7.4, 4.9])
heading("4.1  业务流程节点颜色说明", 2)
para("每个工作流实例按“当前所处角色”区分节点颜色：", space_after=2)
bullet("自己的待处理节点：蓝色（主动态）", bold_prefix="🔵 ")
bullet("他人已完成节点：绿灰色（done）", bold_prefix="🟢 ")
bullet("未处理节点：灰色（pending）", bold_prefix="⚪ ")
bullet("当前节点：琥珀色高亮（active/current）", bold_prefix="🟠 ")
note("每个角色登录看到的同一工作流节点颜色不同，取决于该节点是否属于自己，这是系统的核心设计。")
heading("4.2  工作台待办数字与独立页面一致", 2)
para("工作台“我的待办”数量、审批中心、以及“我的待办”独立列表页，三处均来自同一套待办口径，保证数字一致。ADMIN/GM 角色可见全部待办，其他角色仅见分配给自己的任务。")

# ================= 5 业务模块 =================
heading("5  业务模块详细操作", 1)

heading("5.1  客户管理", 2)
para("维护客户档案与结算信息。入口：工作台 → 客户档案。")
make_table(["操作", "步骤", "关键字段"], [
    ["新增客户", "点击「新增客户」→ 填写表单 → 保存", "编码C001、名称、税号、地址、联系人、电话、行业、结算周期（月结30/60/90/款到发货）、开户行、账号"],
    ["编辑客户", "客户卡片 → 编辑 → 修改 → 保存", "同上"],
    ["搜索", "顶部搜索框输入名称/编码模糊搜索", "—"],
], widths=[3.0, 6.0, 6.5])
note("订单/合同选择客户时，下拉可直接搜索客户名称；客户编码建议与结算主体关联，便于财务对账。")

heading("5.2  销售订单", 2)
para("订单贯穿「销售发起 → 审批 → 生产 → 发货 → 结算」全流程。入口：工作台 → 销售订单。")
make_table(["操作", "步骤"], [
    ["新建订单", "点击「新建订单」→ 选择客户 → 填写公司主体、开票类型（专票/普票/现金）、预收款 → 添加明细（工件名称、规格、计价方式、数量、单位、单价、料属、涂料规格）→ 保存为草稿"],
    ["提交订单", "草稿状态 → 提交 → 自动启动核心生产流（销售→部门主管→财务→总经理→仓管→运营→财务→厂长→质检→归档）"],
    ["生效/退单", "流程完成后订单自动变为「已生效」；异常可退单"],
    ["下加工单", "已生效订单 → 生成加工工单"],
    ["调价申请", "订单 → 发起调价（涨价/降价、固定金额/百分比、原因）"],
    ["利润分析", "查看单订单成本与利润明细"],
], widths=[3.2, 12.3])
make_table(["订单状态", "说明"], [
    ["草稿 DRAFT", "已保存未提交，可编辑"],
    ["待生效 SUBMITTED", "已提交，等待流程审批完成"],
    ["已生效 EFFECTIVE", "流程全部通过，可下加工单"],
    ["生产中 PROCESSING", "已有工单在车间流转"],
    ["待发货 PENDING_DELIVERY", "生产完成待发货"],
    ["已发货 DELIVERED", "已发货登记"],
    ["已结算 CLOSED", "收款核销完成"],
    ["已退单/取消", "异常或取消"],
], widths=[4.5, 11.0])
note("销售订单的客户名称仅“负责销售本人 + 总经理”可见完整名称，其余角色自动打码缩写（数据隐私保护）。")

heading("5.3  加工工单", 2)
para("管理车间加工任务与进度。入口：工作台 → 加工工单。")
make_table(["操作", "步骤"], [
    ["新建工单", "选择来源订单 → 填写批次号、车间、计划数量、计划交期 → 保存"],
    ["下达工单", "工单「下达」→ 状态变为已下达，车间开始排产"],
    ["填完工", "生产完成后「填完工」→ 录入完工单"],
    ["成本查看", "查看单工单人工/材料/制造费用"],
    ["打印工艺单", "打印工艺单 / 质检单"],
], widths=[3.2, 12.3])
make_table(["工单状态", "说明"], [
    ["CREATED 已创建", "刚生成，未下达"],
    ["RELEASED 已下达", "已下达车间"],
    ["PROCESSING 生产中", "车间进行中"],
    ["COMPLETED 已完工", "已报完工"],
    ["CONFIRMED 已确认", "完工单确认完成"],
], widths=[4.5, 11.0])

heading("5.4  完工单", 2)
para("记录加工完成情况与成本核算。入口：工作台 → 完工单。")
make_table(["操作", "步骤"], [
    ["录入完工", "选择加工单 → 填写完工数、合格数、返工数、废品数、工时、人工费、制造费、涂料用量 → 保存草稿"],
    ["确认完工", "提交确认 → 自动生成完工单流程（厂长发起→质检确认→运营归档）"],
    ["打印质检单", "打印质检单留档"],
], widths=[3.2, 12.3])

heading("5.5  领料出库", 2)
para("物料出库管理。入口：工作台 → 领料出库。")
make_table(["操作", "步骤"], [
    ["生成领料单", "领料单由加工单自动生成（或手动发起）"],
    ["确认出库", "仓管核对物料明细后「确认出库」→ 库存自动扣减"],
    ["拒领", "物料异常可「拒领」退回"],
    ["打印领料单", "打印领料单留档"],
], widths=[3.2, 12.3])
make_table(["领料单状态", "说明"], [["PENDING 待处理", "已生成待仓管确认"], ["CONFIRMED 已出库", "已确认出库"], ["REJECTED 已拒领", "已拒绝领料"]], widths=[4.5, 11.0])

heading("5.6  库存管理", 2)
para("实时查看物料库存。入口：工作台 → 库存。")
make_table(["操作", "步骤"], [
    ["库存查询", "按物料类别（如涂料/辅料）、关键字筛选，查看库存量与安全库存"],
    ["新增物料", "填写物料编码、名称、规格、单位、分类、库位、期初库存、安全库存、单价"],
    ["低库存预警", "库存量 ≤ 安全库存自动标记预警（工作台 KPI 同步显示低库存告警数）"],
], widths=[3.2, 12.3])

heading("5.7  财务单据", 2)
para("财务收支与应收应付管理。入口：工作台 → 财务单据 / 应收管理。")
make_table(["操作", "步骤"], [
    ["登记收款", "选择待收款单据 → 登记收款（区分专票/普票/现金）"],
    ["单据查看", "按单据类型（应收/应付/收款/付款）与状态筛选"],
    ["核销", "收款到账后核销应收（未核销金额自动计算）"],
], widths=[3.2, 12.3])
make_table(["财务单据状态", "说明"], [["DRAFT 草稿", "自动生成未结算"], ["OPEN 待结算", "待收款/付款"], ["SETTLED 已结算", "已完成核销"], ["CANCELLED 已取消", "已作废"]], widths=[4.5, 11.0])

heading("5.8  采购订单与采购申请", 2)
para("采购全流程管理。入口：工作台 → 采购订单 / 采购申请。")
make_table(["模块", "操作", "状态"], [
    ["采购申请 PR", "发起申请 → 部门主管审批 → 财务审核 → 总经理终审 → 转为采购订单", "流程驱动"],
    ["采购订单 PO", "创建采购订单 → 下单 → 收货入库（入库后库存自动增加）", "DRAFT/ORDERED/RECEIVED/CANCELLED"],
], widths=[3.2, 8.6, 3.7])

heading("5.9  工资管理", 2)
para("工资表生成与发放管理。入口：工作台 → 工资管理（财务角色可见）。支持按员工生成工资单、批量发放并标记发放状态。")

heading("5.10  调价申请", 2)
para("合同价与实际到款有差异时，销售发起调价。入口：工作台 → 调价申请。")
make_table(["步骤", "操作"], [
    ["1 发起", "选择已生效订单 → 选择调价类型（涨价/降价）、调价方式（固定金额/百分比）、输入金额与原因"],
    ["2 审批", "总经理审批（必须销售发起 + 总经理通过才能生效）"],
    ["3 联动", "审批通过后财务到账登记联动差额"],
], widths=[2.5, 13.0])

# ================= 6 业务流程 =================
heading("6  业务流程操作指引", 1)
para("系统内置 7 条标准业务流程。发起后在工作台「业务流程」与审批中心可见并推进。")

heading("6.1  核心生产流（12 节点 · 销售发起）", 2)
flow_box("核心生产流", ["销售发起", "部门主管审批", "财务审核", "总经理审批", "仓管来货登记", "运营核对", "财务入账", "生产下达", "车间生产", "完工确认", "质检确认", "运营归档"])
para("销售创建销售订单即自动生成该流程实例，节点依次流转；任一节点驳回则退回上一环节。", space_after=2)

heading("6.2  来货登记流程", 2)
flow_box("来货登记", ["仓管发起", "运营核对", "财务入账", "归档"])
para("全部节点通过后，关联的销售订单自动变为「已生效」。", space_after=2)

heading("6.3  费用报销审批", 2)
flow_box("费用报销", ["员工发起", "部门主管初审", "财务审核", "总经理终审(>5000元)"])
para("5000 元以下自动生效；超过 5000 元自动流转到总经理终审。", space_after=2)

heading("6.4  采购请求审批", 2)
flow_box("采购请求", ["采购发起", "部门主管审批", "财务审核", "总经理审批"])
heading("6.5  采购审批流", 2)
flow_box("采购审批", ["采购发起", "部门主管审批", "财务审核", "总经理终审"])
heading("6.6  调价申请审批", 2)
flow_box("调价申请", ["销售发起", "总经理审批"])
heading("6.7  完工单确认", 2)
flow_box("完工确认", ["厂长发起", "质检确认", "运营归档"])

# ================= 7 审批中心 =================
heading("7  审批中心", 1)
para("统一处理当前用户所有待审批事项。入口：左侧导航「审批中心」或工作台待办跳转。")
make_table(["操作", "说明"], [
    ["查看待办", "列出 PENDING 状态的审批任务，显示单据类型、单据号、业务标题、节点、等待时长"],
    ["通过", "点击「通过」→ 可填写意见 → 确认后流程推进至下一节点"],
    ["驳回", "点击「驳回」→ 填写原因 → 流程退回上一节点，业务单据状态同步回退"],
    ["转办", "将任务转交给指定用户处理"],
    ["催办", "对长期未处理的任务发起催办通知"],
], widths=[3.0, 12.5])
note("审批历史会展示各节点的处理人、处理时间与审批意见，供全流程追溯。")

# ================= 8 经营分析 =================
heading("8  经营分析与 AI 助手", 1)
para("经营分析模块（仅 ADMIN/GM/FINANCE 可见）提供三大能力：")
heading("8.1  KPI 仪表盘", 2)
para("展示核心经营指标：今日订单、本月销售额、待审批数、低库存告警等，一屏掌握经营概况。")
heading("8.2  多维透视分析（Pivot）", 2)
para("通过 6 个下拉参数自由组合，无需写 SQL 即可完成任意维度交叉分析：")
make_table(["参数", "说明", "示例"], [
    ["数据源", "选择数据集（订单、客户、采购、财务等）", "订单"],
    ["行维度", "按哪个字段分组（客户、销售、产品、月份…）", "销售员"],
    ["列维度", "交叉列字段（可选）", "月份"],
    ["指标", "要统计的数字（金额、数量…）", "订单金额"],
    ["聚合方式", "求和/平均值/计数/最大值/最小值", "求和"],
    ["图表类型", "柱状图/折线图/饼图/表格", "柱状图"],
], widths=[3.2, 6.8, 5.5])
para("透视表支持点击维度值下钻（drill-down）查看明细记录；支持日期、枚举、数值区间智能筛选；外键维度自动显示名称而非 ID。")
heading("8.3  AI 问答（自然语言）", 2)
para("输入自然语言提问（如“本月各销售员订单金额排名”），系统通过 DeepSeek 双模型解析意图 → 后端真实跑 SQL 聚合 → 生成图文报告。所有数字均来自真实数据，禁止 LLM 编造。")
note("AI 功能需正确配置 DEEPSEEK_API_KEY 等环境变量；意图解析失败时自动回退关键词匹配，不会白屏。")

# ================= 9 流程设计器 =================
heading("9  流程设计器", 1)
para("可视化自定义业务流程。入口：左侧栏「流程设计」或工作台每个工作流标题右侧「编辑」。")
make_table(["操作", "方式"], [
    ["添加节点", "从左侧节点库拖拽到画布"],
    ["连接节点", "鼠标悬停节点左侧小圆 → 拖到下一节点右侧小圆"],
    ["编辑节点", "双击节点 → 弹出属性面板（名称/类型/审批角色/分支条件）"],
    ["删除节点/连线", "右键节点/连线 → 弹出「删除」菜单确认"],
    ["保存发布", "右上「保存」→ 定义持久化，立即可被审批中心调用"],
], widths=[4.0, 11.5])
make_table(["节点类型", "说明"], [
    ["process 业务动作", "发起/提交/发货/归档等非审批节点"],
    ["approve 审批节点", "指定 approver_role（ADMIN/SALES 等），该角色用户待办出现数字徽章"],
    ["fork 分支节点", "基于金额/字段条件分裂多条路径"],
    ["cc 抄送节点", "抄送通知指定角色"],
], widths=[4.0, 11.5])
note("删除工作流定义会连带停掉在跑实例（软停，不丢历史数据），请谨慎操作。")

# ================= 10 管理员后台 =================
heading("10  管理员后台", 1)
heading("10.1  用户管理", 2)
make_table(["动作", "说明"], [
    ["新增用户", "填账号/姓名/角色/手机号；默认密码 123456"],
    ["修改", "可改角色/姓名/手机号/状态"],
    ["重置密码", "单独按钮 → 重置成 123456 并弹提示"],
    ["启停", "状态 ACTIVE ↔ INACTIVE；停用后立即无法登录"],
    ["删除", "软删除（保留历史）；名下有在跑任务则禁止删除"],
], widths=[3.5, 12.0])
heading("10.2  角色管理", 2)
make_table(["动作", "说明"], [
    ["新增角色", "code（英文常量）+ 中文名"],
    ["修改", "可改中文名 + 可见模块权限（pages 字段）"],
    ["删除", "角色下还有用户关联则禁止删除，需先转移用户"],
], widths=[3.5, 12.0])
note("系统启动时按内置默认值强制同步各角色的 pages 权限（GM/ADMIN 为全权限 '*'），保证各角色入口与权限一致。")

# ================= 11 生产大屏 =================
heading("11  生产大屏", 1)
para("面向车间管理者的可视化数据看板（厂长角色可见），展示：")
bullet("各车间产量分布")
bullet("订单状态分布")
bullet("经营快报（今日订单、待审批、低库存等）")
bullet("最近进行中的工单")

# ================= 12 常见问题 =================
heading("12  常见问题（FAQ）", 1)
make_table(["问题", "解答"], [
    ["页面空白/黑屏？", "按 Ctrl+F5 强制刷新浏览器缓存（前端资源带版本号）"],
    ["看不到某些菜单？", "当前账号角色无权限，联系管理员调整角色或 pages 权限"],
    ["登录提示“登录已过期”？", "首次访问新域名无缓存登录态导致，直接输入账号密码登录即可"],
    ["调价申请提交失败？", "确认订单状态为“已生效”且当前用户为销售角色"],
    ["待办数量与审批中心不一致？", "三处共用同一待办口径（ADMIN/GM 看全部，其他角色看自己），如仍不一致请刷新页面"],
    ["如何创建新账号？", "管理员登录 → 系统管理 → 用户管理 → 新增用户"],
    ["AI 问答报“DEEPSEEK_API_KEY 未配置”？", "在环境变量配置 DEEPSEEK_API_KEY 后重启服务"],
    ["时间显示不准确？", "系统统一使用北京时间（UTC+8），刷新页面即可"],
], widths=[4.5, 11.0])

# ================= 13 附录 =================
heading("13  附录", 1)
heading("13.1  本地启动开发环境", 2)
p = doc.add_paragraph()
r = p.add_run("pip install -r requirements.txt\npython -m uvicorn app.main:app --host 0.0.0.0 --port 8000\n浏览器打开 http://localhost:8000 → admin / admin123")
set_cn(r, size=9.5, color=GRAY)
heading("13.2  项目关键文件结构", 2)
make_table(["路径", "说明"], [
    ["app/main.py", "FastAPI 入口、启动种子数据（角色/用户/流程定义）、/health 探活短路"],
    ["app/config.py", "全部环境变量（含 DeepSeek 四项）"],
    ["app/core/", "auth.py(JWT)、db.py(SQLAlchemy)、llm.py(DeepSeek客户端)、permissions.py(权限)、pivot.py(透视引擎)"],
    ["app/api/", "auth/workbench/orders/approvals/analysis/ai_analysis/admin_management 等业务接口"],
    ["app/models/", "SQLAlchemy ORM 模型"],
    ["static/", "index.html、app.js(Vue3前端)、style.css(暗色主题)、icons.js(Heroicons)"],
    ["scripts/seed_data.py", "初始化建表 + 角色 + 流程定义 + 示例数据"],
], widths=[5.0, 10.5])
heading("13.3  环境变量说明", 2)
make_table(["变量", "示例值", "说明"], [
    ["PORT", "8000", "服务端口（数字，勿写 ${} 占位）"],
    ["JWT_SECRET", "长随机字符串", "JWT 签名密钥"],
    ["JWT_TTL_MINUTES", "1440", "令牌有效期（分钟）"],
    ["DB_DRIVER / DB_URL", "sqlite / sqlite:////app/data/erp.db", "数据库驱动与连接串"],
    ["DEEPSEEK_API_KEY", "sk-xxx", "DeepSeek 密钥（AI 模块必需）"],
    ["DEEPSEEK_BASE_URL", "https://api.deepseek.com", "API 基址（自动智能拼接 /chat/completions）"],
    ["DEEPSEEK_MODEL_FAST / PRO", "deepseek-v4-flash / deepseek-v4-pro", "意图解析与报告生成模型"],
], widths=[4.5, 5.5, 5.5], font_size=9)
heading("13.4  数据库", 2)
para("系统默认使用 SQLite（data/erp.db，WAL 模式），单文件部署、零运维成本；后续可平滑切换 PostgreSQL（DB_DRIVER=postgres + DB_URL 调整）。")

# ============ 页脚 ============
sec = doc.sections[0]
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("峰业精密ERP · 系统功能及详细操作手册 · v1.0")
set_cn(fr, size=8, color=GRAY)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "峰业精密ERP-系统功能及详细操作手册.docx")
doc.save(out)
print("已生成:", out)
